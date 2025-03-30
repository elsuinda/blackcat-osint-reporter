from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import pdfkit
from datetime import datetime
import os

def generate_word_report(report_data, output_path):
    try:
        doc = Document()
        
        # Add styles
        styles = doc.styles
        title_style = styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle_style = styles.add_style('SubtitleStyle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_style.font.size = Pt(14)
        subtitle_style.font.bold = True
        
        normal_style = styles.add_style('NormalStyle', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.size = Pt(12)
        
        # Add title
        doc.add_paragraph('OSINT REPORT', style='TitleStyle')
        doc.add_paragraph(f"Report Number: {report_data['report_number']}", style='SubtitleStyle')
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style='NormalStyle')
        doc.add_paragraph('', style='NormalStyle')  # Empty line
        
        # Add report details
        doc.add_paragraph('Source:', style='SubtitleStyle')
        doc.add_paragraph(report_data['source'], style='NormalStyle')
        
        if report_data.get('title'):
            doc.add_paragraph('Title:', style='SubtitleStyle')
            doc.add_paragraph(report_data['title'], style='NormalStyle')
        
        doc.add_paragraph('Post Date:', style='SubtitleStyle')
        doc.add_paragraph(report_data['post_date'], style='NormalStyle')
        
        doc.add_paragraph('Summary:', style='SubtitleStyle')
        doc.add_paragraph(report_data['summary'], style='NormalStyle')
        
        if report_data.get('url'):
            doc.add_paragraph('URL:', style='SubtitleStyle')
            doc.add_paragraph(report_data['url'], style='NormalStyle')
        
        if report_data.get('user_id'):
            doc.add_paragraph('User ID:', style='SubtitleStyle')
            doc.add_paragraph(report_data['user_id'], style='NormalStyle')
        
        # Add image if exists
        if report_data.get('image_path') and os.path.exists(report_data['image_path']):
            doc.add_paragraph('Image:', style='SubtitleStyle')
            doc.add_picture(report_data['image_path'], width=Inches(4.0))
        
        # Save document
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error generating Word report: {e}")
        return False

def generate_pdf_report(report_data, output_path):
    try:
        # First generate HTML
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>OSINT Report {report_data['report_number']}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 2cm; }}
                h1 {{ color: #2c3e50; text-align: center; }}
                h2 {{ color: #3498db; border-bottom: 1px solid #eee; padding-bottom: 5px; }}
                .report-info {{ margin-bottom: 20px; }}
                .section {{ margin-bottom: 15px; }}
            </style>
        </head>
        <body>
            <h1>OSINT REPORT</h1>
            <div class="report-info">
                <p><strong>Report Number:</strong> {report_data['report_number']}</p>
                <p><strong>Date:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            </div>
            
            <div class="section">
                <h2>Source</h2>
                <p>{report_data['source']}</p>
            </div>
            
            {f'<div class="section"><h2>Title</h2><p>{report_data["title"]}</p></div>' if report_data.get('title') else ''}
            
            <div class="section">
                <h2>Post Date</h2>
                <p>{report_data['post_date']}</p>
            </div>
            
            <div class="section">
                <h2>Summary</h2>
                <p>{report_data['summary']}</p>
            </div>
            
            {f'<div class="section"><h2>URL</h2><p>{report_data["url"]}</p></div>' if report_data.get('url') else ''}
            
            {f'<div class="section"><h2>User ID</h2><p>{report_data["user_id"]}</p></div>' if report_data.get('user_id') else ''}
        </body>
        </html>
        """
        
        # Generate PDF from HTML
        options = {
            'page-size': 'A4',
            'margin-top': '1.5cm',
            'margin-right': '1.5cm',
            'margin-bottom': '1.5cm',
            'margin-left': '1.5cm',
            'encoding': 'UTF-8',
            'quiet': ''
        }
        
        pdfkit.from_string(html_content, output_path, options=options)
        return True
    except Exception as e:
        print(f"Error generating PDF report: {e}")
        return False
